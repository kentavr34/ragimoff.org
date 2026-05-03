#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Обновляет DOCX: применяет все накопленные текстовые правки из HTML.
Для циклических замен (A→B, B→C, C→A) использует временные метки.
Сохраняет как KLINIK_PSIXIATRIYA_2026.docx на рабочем столе.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

SRC = r"D:\BOOKS\FINAL\kp_new\KLINIK_PSIXIATRIYA.docx"
DST = r"C:\Users\SAM\Desktop\KLINIK_PSIXIATRIYA_2026.docx"

# ═══════════════════════════════════════════════════════════════════════
# Таблица замен — ПОРЯДОК КРИТИЧЕН.
# Циклические (A→B→C→A) делаем через temp-метки (§xx§).
# ═══════════════════════════════════════════════════════════════════════

REPLACEMENTS = [

    # ── bolme-11: BƏDƏNSƏL DISSTRES ───────────────────────────────────
    ("BƏDƏNSƏL DISSTRES VƏ ƏLAQƏLİ POZUNTULAR",
     "BƏDƏN DİSSTRESİ VƏ HİSSİYƏTİ POZUNTULARI"),
    ("Bədənsəl disstres və əlaqəli pozuntular",
     "Bədən distresi və hissiyəti pozuntuları"),
    ("BƏDƏNSƏL DISSTRES POZUNTUSU",
     "BƏDƏN DİSSTRES POZUNTUSU"),
    ("Bədənsəl disstres pozuntusu",
     "Bədən disstres pozuntusu"),
    ("BƏDƏN BÜTÖVLÜYÜ DİSFORİYASI",
     "BƏDƏN BÜTÖVLÜYÜNÜ QAVRAMA POZUNTUSU"),
    ("Bədən Bütövlüyü Disforiyasi",
     "Bədən Bütövlüyünü Qavrama Pozuntusu"),
    ("Bədən bütövlüyü disforiyasi",
     "Bədən bütövlüyünü qavrama pozuntusu"),

    # ── bolme-01: ICD кода (циклические — через §temp§) ───────────────
    # Оригинал: learning=F82, coordination=F90, ADHD=F81
    # Нужно:    learning=F81, coordination=F82, ADHD=F90
    ("F82 — Hərəkət koordinasiyasının spesifik pozuntusu", "§01FA§"),
    ("F90 — Hiperkinetik pozuntular",                      "§01FB§"),
    ("F81 — Məktəb bacarıqlarının spesifik pozuntuları",   "§01FC§"),
    ("§01FA§", "F81 — Məktəb bacarıqlarının spesifik pozuntuları"),
    ("§01FB§", "F82 — Hərəkət koordinasiyasının spesifik pozuntusu"),
    ("§01FC§", "F90 — Hiperkinetik pozuntular"),

    # DSM для bolme-01 (Title Case — именно так в DOCX):
    # Оригинал: learning=315.4 DCD, coordination=314.01 ADHD, ADHD=315.00 Learning
    # Нужно:    learning=315.00 Learning, coordination=315.4 DCD, ADHD=314.01 ADHD
    ("315.4 — Developmental Coordination Disorder",           "§01DA§"),
    ("314.01 — Attention-Deficit/Hyperactivity Disorder",     "§01DB§"),
    ("315.00 — Specific Learning Disorder",                   "§01DC§"),
    ("§01DA§", "315.00 — Specific Learning Disorder"),
    ("§01DB§", "315.4 — Developmental Coordination Disorder"),
    ("§01DC§", "314.01 — Attention-Deficit/Hyperactivity Disorder"),

    # Заголовки расстройств bolme-01 (UPPERCASE + mixed, + с тире в XBT-11 строке)
    ("6A04 İNKİŞAF ÖYRƏNMƏ POZUNTUSU", "6A03 İNKİŞAF ÖYRƏNMƏ POZUNTUSU"),
    # Строка XBT-11 в комбинированном абзаце (с тире, sentence case из DOCX)
    ("6A04 — İnkişaf öyrənmə pozuntusu", "6A03 — İnkişaf öyrənmə pozuntusu"),
    ("6A05 İNKİŞAF HƏRƏKİ KOORDİNASİYA POZUNTUSU",
     "6A04 İNKİŞAF HƏRƏKİ KOORDİNASİYA POZUNTUSU"),
    ("6A05 — İnkişaf hərəki koordinasiya pozuntusu",
     "6A04 — İnkişaf hərəki koordinasiya pozuntusu"),
    ("6A03 DİQQƏT DEFİSİTİ VƏ HİPERAKTİVLİK POZUNTUSU",
     "6A05 DİQQƏT DEFİSİTİ VƏ HİPERAKTİVLİK POZUNTUSU"),
    # ADHD inline (DOCX title case + dash)
    ("6A03 — Diqqət Defisiti Və Hiperaktivlik Pozuntusu (Ddhp)",
     "6A05 — Diqqət Defisiti Və Hiperaktivlik Pozuntusu (Ddhp)"),

    # ── bolme-13: перепутанные F-коды (циклические) ───────────────────
    # Оригинал: 6C72=F52.7, 6C71=F63.1, 6C73=F63.8, 6C70=F63.2
    # Нужно:    6C72=F63.1, 6C71=F63.2, 6C73=F52.7, 6C70=F63.8
    ("F52.7 — Həddindən artıq cinsi əlaqə",         "§13FA§"),
    ("F63.1 — Patoloji oduvurma",                   "§13FB§"),
    ("F63.8 — İmpuls nəzarəti pozuntusu",            "§13FC§"),
    ("F63.2 — Patoloji oğurlama",                   "§13FD§"),
    ("§13FA§", "F63.1 — Patoloji oduvurma"),
    ("§13FB§", "F63.2 — Patoloji oğurlama"),
    ("§13FC§", "F52.7 — Həddindən artıq cinsi əlaqə"),
    ("§13FD§", "F63.8 — Digər vərdiş və impuls pozuntuları"),

    # DSM для bolme-13 (Title Case — именно так в DOCX):
    # Оригинал: 6C72=Compulsive Sexual Behavior Disorder (wrong),
    #           6C71=312.33 Pyromania (wrong), 6C73=312.34 IED (wrong), 6C70=312.32 Kleptomania (wrong)
    # Нужно:    6C72=312.33 Pyromania, 6C71=312.32 Kleptomania,
    #           6C73=— — Compulsive sexual behaviour disorder, 6C70=312.34 IED
    ("— — Compulsive Sexual Behavior Disorder",      "§13DA§"),
    ("— — Compulsive sexual behavior disorder",      "§13DA§"),
    ("— — compulsive sexual behavior disorder",      "§13DA§"),
    ("312.33 — Pyromania",                           "§13DB§"),
    ("312.34 — Intermittent Explosive Disorder",     "§13DC§"),
    ("312.34 — Intermittent explosive disorder",     "§13DC§"),
    ("312.32 — Kleptomania",                         "§13DD§"),
    ("§13DA§", "312.33 — Pyromania"),
    ("§13DB§", "312.32 — Kleptomania"),
    ("§13DC§", "— — Compulsive sexual behaviour disorder"),
    ("§13DD§", "312.34 — Intermittent Explosive Disorder"),

    # ── "Davam edən" → "Persistentli" ────────────────────────────────
    ("Davam edən sayıqlama pozuntusu",  "Persistentli sayıqlama pozuntusu"),
    ("DAVAM EDƏN SAYIQLAMA POZUNTUSU", "PERSİSTENTLİ SAYIQLAMA POZUNTUSU"),
    ("Davam edən halüsinasiya pozuntusu",  "Persistentli halüsinasiya pozuntusu"),
    ("DAVAM EDƏN HALÜSİNASİYA POZUNTUSU","PERSİSTENTLİ HALÜSİNASİYA POZUNTUSU"),
    ("Davam edən əhval pozuntusu",         "Persistentli əhval pozuntusu"),
    ("DAVAM EDƏN ƏHVAL POZUNTUSU",         "PERSİSTENTLİ ƏHVAL POZUNTUSU"),

    # ── İnsomnia → Yuxusuzluq ─────────────────────────────────────────
    ("İnsomnia pozuntusu",   "Yuxusuzluq pozuntusu (İnsomniya)"),
    ("INSOMNIA POZUNTUSU",   "YUXUSUZLUQ POZUNTUSU (İNSOMNİYA)"),
    ("İNSOMNİA POZUNTUSU",   "YUXUSUZLUQ POZUNTUSU (İNSOMNİYA)"),
]


# ── Функции работы с документом ──────────────────────────────────────────────

def para_fulltext(para):
    return ''.join(r.text for r in para.runs)

def replace_in_para(para, old, new):
    full = para_fulltext(para)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not para.runs:
        return False
    # Кладём весь текст в первый run, остальные обнуляем
    para.runs[0].text = new_full
    for r in para.runs[1:]:
        r.text = ''
    return True

def replace_in_table(table, old, new):
    n = 0
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if replace_in_para(p, old, new):
                    n += 1
    return n

def apply_replacement(doc, old, new):
    n = 0
    for p in doc.paragraphs:
        if replace_in_para(p, old, new):
            n += 1
    for t in doc.tables:
        n += replace_in_table(t, old, new)
    return n


# ── Главный процесс ──────────────────────────────────────────────────────────

print(f"Открываем: {SRC}")
doc = Document(SRC)

total = 0
for old, new in REPLACEMENTS:
    n = apply_replacement(doc, old, new)
    if n:
        total += n
        old_s = old[:55].replace('\n', ' ')
        new_s = new[:55].replace('\n', ' ')
        print(f"  [{n:2d}] {old_s!r} → {new_s!r}")

print(f"\nВсего замен: {total}")
print(f"Сохраняем: {DST}")
doc.save(DST)
print("Готово.")
