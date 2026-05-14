#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build KLINIK_PSIXIATRIYA_6.docx from updated site HTML."""
import re
import sys
import io
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

ROOT = Path(__file__).parent.resolve()
SITE = ROOT / "klinik-psixiatriya"
BUILD_DIR = ROOT / "_build"
BUILD_DIR.mkdir(exist_ok=True)

OUT_HTML = BUILD_DIR / "book_combined.html"
OUT_DOCX = SITE / "KLINIK_PSIXIATRIYA_6.docx"
REFERENCE_DOCX = SITE / "KLINIK_PSIXIATRIYA_5.docx"

ORDER = [
    "mugeddime.html",
    "giris-yekun.html",
    "01-6A0-neyroinkisaf.html",
    "02-6A2-sizofreniya-spektri.html",
    "03-6A4-katatoniya.html",
    "04-6A6-ehval-pozuntulari.html",
    "05-6B0-narahatliq.html",
    "06-6B2-okp.html",
    "07-6B4-stress.html",
    "08-6B6-dissosiativ.html",
    "09-6B8-qida-qebulu.html",
    "10-6C0-ifrazat.html",
    "11-6C2-bedensel-disstres.html",
    "12-6C4-madde-asililiq.html",
    "13-6C7-impuls-nezareti.html",
    "14-6C9-pozucu-davranis.html",
    "15-6D1-sexsiyyet.html",
    "16-6D3-parafilik.html",
    "17-6D5-faktitioz.html",
    "18-6D7-neyrokoqnitiv.html",
    "19-6E2-perinatal.html",
    "20-6E4-psixosomatik.html",
    "21-6E6-ikincili.html",
    "22-7AB-yuxu.html",
    "23-HA-cinsi-saglamliq.html",
    "elave-acde.html",
    "elave-skalalar.html",
    "yekun.html",
]

MAIN_RE = re.compile(r'<main\b[^>]*>(.*?)</main>', re.DOTALL | re.IGNORECASE)
SUPP_MARKER_RE = re.compile(r'<!--\s*/?BOOK-SUPPLEMENT:[^>]*-->', re.IGNORECASE)


def extract_main(html: str) -> str:
    m = MAIN_RE.search(html)
    if not m:
        return ""
    body = m.group(1)
    body = SUPP_MARKER_RE.sub("", body)
    body = re.sub(r'<aside\b.*?</aside>', '', body, flags=re.DOTALL | re.IGNORECASE)
    body = re.sub(r'<script\b.*?</script>', '', body, flags=re.DOTALL | re.IGNORECASE)
    body = re.sub(r'<style\b.*?</style>', '', body, flags=re.DOTALL | re.IGNORECASE)
    # Strip nav blocks (page-nav at the bottom, breadcrumbs, etc.)
    body = re.sub(r'<nav\b.*?</nav>', '', body, flags=re.DOTALL | re.IGNORECASE)
    return body.strip()


def strip_links(html: str) -> str:
    """Convert <a href="...">text</a> → text (book has no clickable links)."""
    # Strip ALL <a> wrappers, keep inner text
    html = re.sub(r'<a\b[^>]*>(.*?)</a>', r'\1', html, flags=re.DOTALL | re.IGNORECASE)
    return html


def strip_inline_bold(html: str) -> str:
    """Remove <strong>/<b> inside paragraphs and list items (AI-style label markers).
    Preserve bold inside table cells (<th>/<td>)."""
    # Process paragraphs and li only, leaving table cells intact.
    def strip_in(tag, text):
        pattern = re.compile(rf'(<{tag}\b[^>]*>)(.*?)(</{tag}>)', re.DOTALL | re.IGNORECASE)
        def repl(m):
            inner = m.group(2)
            inner = re.sub(r'</?(?:strong|b)\b[^>]*>', '', inner, flags=re.IGNORECASE)
            return m.group(1) + inner + m.group(3)
        return pattern.sub(repl, text)

    for tag in ('p', 'li', 'em'):
        html = strip_in(tag, html)
    return html


def shift_levels(html: str) -> str:
    """Chapter h2 → h1; everything else demoted 1 level."""
    def chapter_to_h1(m):
        attrs = m.group(1)
        if 'data-chapter' not in attrs:
            attrs += ' data-chapter="1"'
        return f'<h1{attrs}>{m.group(2)}</h1>'
    html2, n = re.subn(r'<h2([^>]*)>(.*?)</h2>', chapter_to_h1, html, count=1, flags=re.DOTALL)
    if n == 0:
        return html

    out = []
    i = 0
    pat = re.compile(r'<(/?)h([1-5])(\b[^>]*)>', re.IGNORECASE)
    for m in pat.finditer(html2):
        out.append(html2[i:m.start()])
        slash, lvl, attrs = m.group(1), int(m.group(2)), m.group(3)
        if lvl == 1 and 'data-chapter' in attrs:
            out.append(f'<{slash}h1{attrs}>')
        else:
            new_lvl = min(lvl + 1, 6)
            out.append(f'<{slash}h{new_lvl}{attrs}>')
        i = m.end()
    out.append(html2[i:])
    return "".join(out)


def split_yekun_signature(html: str) -> str:
    """Insert a page-break marker before the author's closing signature in yekun."""
    # Find the closing block (Kitabın məqsədi → signature)
    marker = '<p><strong>Kitabın məqsədi:</strong>'
    if marker in html:
        # Replace preceding <hr> with page-break div, drop <hr>
        html = re.sub(r'<hr\s*/?>\s*' + re.escape(marker),
                      '<div class="pagebreak-before"></div>' + marker,
                      html, flags=re.IGNORECASE)
    return html


def build_title_page() -> str:
    """No HTML title page — pandoc generates one from metadata (Title style,
    not Heading 1, so it stays out of the TOC)."""
    return ""


def main():
    print("Extracting content...")
    parts = [build_title_page()]
    for fn in ORDER:
        path = SITE / fn
        if not path.exists():
            print(f"  SKIP missing: {fn}")
            continue
        html = path.read_text(encoding="utf-8")
        body = extract_main(html)
        if not body:
            continue
        is_chapter = re.match(r'^\d{2}-', fn) is not None
        if is_chapter:
            body = shift_levels(body)
        if fn == "yekun.html":
            body = split_yekun_signature(body)
        body = strip_links(body)
        body = strip_inline_bold(body)
        parts.append(f"<!-- ===== {fn} ===== -->\n{body}\n")
        print(f"  OK {fn}")

    combined = "\n".join(parts)
    full = (
        "<!DOCTYPE html>\n<html lang='az'><head><meta charset='utf-8'>"
        "<title>Klinik Psixiatriya</title></head><body>\n"
        + combined +
        "\n</body></html>"
    )
    OUT_HTML.write_text(full, encoding="utf-8")
    print(f"\nCombined HTML: {OUT_HTML} ({len(full):,} chars)")

    print("\nRunning pandoc...")
    import pypandoc
    # No title/subtitle/author/date metadata — we build a manual title page
    # in fix_docx() so it follows TYPOGRAPHY.md rule #1 (single page: name on
    # top, slogan below, author/city/year at bottom). Pandoc's default title
    # block would have separate Title/Subtitle/Author paragraphs in column.
    extra_args = [
        f"--reference-doc={REFERENCE_DOCX}",
        "--toc",
        "--toc-depth=2",
        "--metadata=toc-title:MÜNDƏRİCAT",
        "--standalone",
    ]
    pypandoc.convert_file(
        str(OUT_HTML),
        "docx",
        outputfile=str(OUT_DOCX),
        extra_args=extra_args,
        format="html",
    )
    print(f"DOCX: {OUT_DOCX}")

    print("\nPost-processing DOCX...")
    fix_docx(OUT_DOCX)
    print(f"\nDONE: {OUT_DOCX}")
    print(f"Size: {OUT_DOCX.stat().st_size / 1024:.0f} KB")


def fix_docx(path: Path):
    """Post-process DOCX per TYPOGRAPHY.md rules:
      1. Manual title page (book name + slogan top, author/city/year bottom)
      2. Heading hierarchy: H1 28pt bold > H2 20pt bold > H3 14pt bold > H4 12pt bold
      3. Page break ONLY before Heading 1 (chapter) and Heading 2 whose text
         starts with an ICD code pattern (disorder title). Sub-sections do NOT
         get page breaks.
      4. Table cell indent + borders normalized.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    doc = Document(str(path))

    # ── A. Configure Heading styles (TYPOGRAPHY.md rule #2) ──
    def style_heading(name, size_pt, bold=True, color=None, align=None):
        try:
            s = doc.styles[name]
        except KeyError:
            return
        s.font.name = "Times New Roman"
        s.font.size = Pt(size_pt)
        s.font.bold = bold
        if color is not None:
            s.font.color.rgb = RGBColor(*color)
        if align is not None:
            s.paragraph_format.alignment = align
        # Reasonable spacing — no huge gaps
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    style_heading("Heading 1", 28, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.CENTER)
    style_heading("Heading 2", 20, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT)
    style_heading("Heading 3", 14, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT)
    style_heading("Heading 4", 12, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT)
    style_heading("Heading 5", 11, bold=True, color=(0,0,0), align=WD_ALIGN_PARAGRAPH.LEFT)
    # Force black on Word's default-blue Heading 6-9 and Hyperlink
    for name in ("Heading 6", "Heading 7", "Heading 8", "Heading 9"):
        try:
            s = doc.styles[name]
            s.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass
    try:
        s = doc.styles["Hyperlink"]
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.font.underline = False  # no underline — looks like body text
    except KeyError:
        pass

    # Tighten Normal/Body spacing + Russian-academic first-line indent
    try:
        n = doc.styles["Normal"]
        n.font.name = "Times New Roman"
        n.font.size = Pt(11)
        n.paragraph_format.space_after = Pt(4)
        n.paragraph_format.space_before = Pt(0)
        n.paragraph_format.line_spacing = 1.25
        n.paragraph_format.first_line_indent = Cm(0.5)
        n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        pass

    # First-paragraph-after-heading should NOT be indented (typical textbook
    # convention). Pandoc tags such paragraphs with 'First Paragraph' style.
    for sty_name in ("First Paragraph", "FirstParagraph"):
        try:
            fp = doc.styles[sty_name]
            fp.font.name = "Times New Roman"
            fp.font.size = Pt(11)
            fp.paragraph_format.first_line_indent = Cm(0)
            fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass

    # ── B. Page headers/footers (TYPOGRAPHY.md rule #6) ──
    # Even (left) page header: chapter title (will be set dynamically by Word
    # via STYLEREF field). Odd (right) page header: book name. Page numbers
    # in outer corners via PAGE field.
    # Enable odd-and-even page headers/footers globally (settings.xml)
    settings_el = doc.settings.element
    if settings_el.find(qn('w:evenAndOddHeaders')) is None:
        settings_el.append(OxmlElement('w:evenAndOddHeaders'))

    for section in doc.sections:
        section.different_first_page_header_footer = True
        # Title page: no header/footer
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        # Even-page (left) header — STYLEREF "Heading 1" pulls chapter title
        eh = section.even_page_header
        eh.is_linked_to_previous = False
        ehp = eh.paragraphs[0]
        ehp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ehp.text = ""
        run = ehp.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.italic = True
        # Append STYLEREF field
        fld_begin = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:ascii'),'Times New Roman'); rFonts.set(qn('w:hAnsi'),'Times New Roman')
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'),'18')
        i_el = OxmlElement('w:i')
        rPr.append(rFonts); rPr.append(sz); rPr.append(i_el)
        fld_begin.append(rPr)
        fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'),'begin')
        fld_begin.append(fb)
        ehp._p.append(fld_begin)
        instr = OxmlElement('w:r')
        instr.append(rPr)
        instr_t = OxmlElement('w:instrText'); instr_t.text = ' STYLEREF "Heading 1" \\* MERGEFORMAT '
        instr_t.set(qn('xml:space'),'preserve')
        instr.append(instr_t)
        ehp._p.append(instr)
        fld_end = OxmlElement('w:r')
        fld_end.append(rPr)
        fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'),'end')
        fld_end.append(fe)
        ehp._p.append(fld_end)
        # Odd-page (right) header — book title
        oh = section.header
        oh.is_linked_to_previous = False
        ohp = oh.paragraphs[0]
        ohp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ohp.text = ""
        r2 = ohp.add_run("KLİNİK PSİXİATRİYA · Diaqnostika və terapiya standartları")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(9)
        r2.italic = True
        # Footer page numbers (both pages, outer corners via centered PAGE field)
        for foot, align in [(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
                            (section.footer, WD_ALIGN_PARAGRAPH.RIGHT)]:
            foot.is_linked_to_previous = False
            fp = foot.paragraphs[0]
            fp.alignment = align
            fp.text = ""
            r3 = fp.add_run()
            r3.font.name = "Times New Roman"
            r3.font.size = Pt(10)
            # PAGE field
            fb = OxmlElement('w:r')
            fbb = OxmlElement('w:fldChar'); fbb.set(qn('w:fldCharType'),'begin')
            fb.append(fbb); fp._p.append(fb)
            it = OxmlElement('w:r')
            itt = OxmlElement('w:instrText'); itt.text = ' PAGE '
            itt.set(qn('xml:space'),'preserve')
            it.append(itt); fp._p.append(it)
            fe = OxmlElement('w:r')
            fee = OxmlElement('w:fldChar'); fee.set(qn('w:fldCharType'),'end')
            fe.append(fee); fp._p.append(fe)

    # ── 1. Force-clear first-line indent in EVERY table cell paragraph ──
    def kill_indent(par):
        pPr = par._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        # Explicit zero overrides style inheritance.
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:start'), '0')
        # Remove any opposing attributes
        for attr in ('w:hanging', 'w:firstLineChars', 'w:leftChars', 'w:startChars'):
            if ind.get(qn(attr)):
                del ind.attrib[qn(attr)]

    def set_cell_border(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tcBorders = tc_pr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tc_pr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = tcBorders.find(qn(f'w:{edge}'))
            if b is None:
                b = OxmlElement(f'w:{edge}')
                tcBorders.append(b)
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:color'), '666666')

    for tbl in doc.tables:
        tbl.autofit = True
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_border(cell)
                for par in cell.paragraphs:
                    kill_indent(par)
                    # Also clear pStyle indent inheritance via direct override
                    for child in par._p.iter():
                        if child.tag.endswith('}pStyle'):
                            # Leave pStyle but our w:ind override above wins
                            pass

    # ── 2. Page-break-before rules (TYPOGRAPHY.md rule #3) ──
    # - Every Heading 1 (chapter) starts on a new page.
    # - Heading 2 starts on a new page ONLY if its text begins with an ICD
    #   code pattern (e.g. "6A00 ", "HA00 ", "7A40 ") — i.e. it's a disorder
    #   title. Chapter-intro Heading 2 (like "Bu bölmənin tərkibi") does NOT
    #   get a page break.
    # - Sub-section headings (Heading 3/4/5) NEVER get page breaks.
    ICD_RE = re.compile(r'^\s*([0-9][A-Z][0-9][0-9A-Z]?|HA[0-9]{2})\b')
    H1_STYLES = {'Heading1', 'Heading 1', 'heading 1'}
    H2_STYLES = {'Heading2', 'Heading 2', 'heading 2'}
    # Strip ANY pre-existing page-break-before from non-target paragraphs.
    SUB_HEADINGS = {'Heading3','Heading 3','heading 3',
                    'Heading4','Heading 4','heading 4',
                    'Heading5','Heading 5','heading 5',
                    'Heading6','Heading 6','heading 6'}
    body = doc.element.body
    paragraphs = body.findall(qn('w:p'))

    def text_of(p):
        return "".join(t.text or "" for t in p.iter(qn('w:t')))

    def add_pbb(p):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        if pPr.find(qn('w:pageBreakBefore')) is None:
            pPr.insert(0, OxmlElement('w:pageBreakBefore'))

    def remove_pbb(p):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            return
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is not None:
            pPr.remove(pbb)

    for p in paragraphs:
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            continue
        sty = pStyle.get(qn('w:val')) or ''
        if sty in H1_STYLES:
            add_pbb(p)
        elif sty in H2_STYLES:
            if ICD_RE.match(text_of(p)):
                add_pbb(p)
            else:
                remove_pbb(p)
        elif sty in SUB_HEADINGS:
            remove_pbb(p)

    # ── 3. Remove first heading's page break (no blank leading page) ──
    for p in paragraphs:
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            continue
        sty = pStyle.get(qn('w:val')) or ''
        if sty in H1_STYLES or sty in H2_STYLES:
            remove_pbb(p)
            break  # only first heading

    # ── 4. Yekun signature page break ──
    for p in paragraphs:
        if text_of(p).startswith("Kitabın məqsədi:"):
            add_pbb(p)
            break

    # ── 4.5 Remove pandoc-auto Title paragraph (we build our own page) ──
    TITLE_STYLES = {'Title', 'title'}
    for p in list(body.findall(qn('w:p'))):
        pPr = p.find(qn('w:pPr'))
        if pPr is None: continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None: continue
        sty = pStyle.get(qn('w:val')) or ''
        if sty in TITLE_STYLES:
            body.remove(p)

    # ── 5. Manual title page (TYPOGRAPHY.md rule #1) ──
    # Insert before the auto-generated TOC. Pandoc puts the TOC right at
    # body start. We prepend: BOOK TITLE (huge, centered) + slogan + author
    # block at bottom.
    title_paragraphs = build_title_page_xml(doc)
    # Insert at the very start of body (before TOC)
    for tp in reversed(title_paragraphs):
        body.insert(0, tp)
    # Add page break to last title-page paragraph so TOC starts on its own page
    if title_paragraphs:
        last = title_paragraphs[-1]
        pPr = last.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            last.insert(0, pPr)
        # Use a run with <w:br w:type="page"/> appended after content
        # Simpler: page break BEFORE the next paragraph (the TOC heading).
        # Find first H1 (TOC heading) after title page and add pbb.
        # Actually TOC heading uses TOC Heading style, not H1. We just rely
        # on pageBreakBefore on the title's own paragraph being absent and
        # add break to the body's first H1.

    # ── 6. Force black on ALL text runs (TYPOGRAPHY.md rule #X: no blue) ──
    # Override any inline color set by pandoc or theme. Exception: title-page
    # grey runs (we preserve their #555555).
    GREY = '555555'
    for p in body.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                continue
            col = rPr.find(qn('w:color'))
            if col is None:
                continue
            val = (col.get(qn('w:val')) or '').upper()
            if val.upper() == GREY.upper():
                continue  # preserve title-page grey
            col.set(qn('w:val'), '000000')
            # Remove any theme-color reference that overrides w:val
            if col.get(qn('w:themeColor')) is not None:
                del col.attrib[qn('w:themeColor')]
            if col.get(qn('w:themeShade')) is not None:
                del col.attrib[qn('w:themeShade')]
            if col.get(qn('w:themeTint')) is not None:
                del col.attrib[qn('w:themeTint')]
    # Same sweep inside table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for r in par._p.findall(qn('w:r')):
                        rPr = r.find(qn('w:rPr'))
                        if rPr is None: continue
                        col = rPr.find(qn('w:color'))
                        if col is None: continue
                        if (col.get(qn('w:val')) or '').upper() == GREY.upper():
                            continue
                        col.set(qn('w:val'), '000000')
                        for attr in ('w:themeColor','w:themeShade','w:themeTint'):
                            if col.get(qn(attr)) is not None:
                                del col.attrib[qn(attr)]

    doc.save(str(path))

    # ── 7. Post-save XML sweep (the python-docx font.color.rgb setter doesn't
    #       strip w:themeColor attribute, so Word still renders Heading 5-9
    #       and Hyperlink in theme-accent blue. We need to edit styles.xml
    #       and theme1.xml directly inside the .docx zip). ──
    _purge_theme_blue(path)


def _purge_theme_blue(path: Path):
    """Force black on every theme-blue source inside the DOCX zip:
       - styles.xml: strip w:themeColor/themeShade/themeTint on every <w:color>;
         set w:val to 000000 (preserve only 555555 grey)
       - theme1.xml: rewrite accent1 + hlink + folHlink to 000000
    """
    import zipfile, shutil, tempfile, re as _re
    tmp = Path(str(path) + ".tmp.zip")
    GREY = '555555'
    NS_W = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'

    def fix_styles(xml: str) -> str:
        # Strip themeColor/themeShade/themeTint from every w:color element;
        # force w:val to 000000 unless it's the preserved grey
        def repl(m):
            tag = m.group(0)
            # Extract current w:val
            mv = _re.search(r'w:val="([0-9A-Fa-f]+)"', tag)
            val = (mv.group(1) if mv else '').upper()
            new_val = val if val == GREY.upper() else '000000'
            # Rebuild a minimal <w:color w:val="..."/>
            self_closing = tag.endswith('/>')
            return f'<w:color w:val="{new_val}"' + ('/>' if self_closing else '>')
        return _re.sub(r'<w:color\b[^/>]*?/?>', repl, xml)

    def fix_theme(xml: str) -> str:
        # Force ALL theme color slots to black (handles `<a:srgbClr val="X"/>`,
        # `<a:srgbClr val="X" />`, and same with whitespace variants).
        return _re.sub(
            r'<a:srgbClr\s+val="[0-9A-Fa-f]+"\s*/>',
            '<a:srgbClr val="000000"/>',
            xml,
        )

    with zipfile.ZipFile(str(path), 'r') as zin, \
         zipfile.ZipFile(str(tmp), 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == 'word/styles.xml':
                data = fix_styles(data.decode('utf-8')).encode('utf-8')
            elif item == 'word/theme/theme1.xml':
                data = fix_theme(data.decode('utf-8')).encode('utf-8')
            zout.writestr(item, data)
    path.unlink()
    tmp.rename(path)


def build_title_page_xml(doc):
    """Build manual title page paragraphs (in document order).
    Returns list of <w:p> elements ready to insert at body start."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def make_paragraph(text, *, size_pt, bold=False, italic=False,
                       align=WD_ALIGN_PARAGRAPH.CENTER,
                       space_before=0, space_after=0, color=(0,0,0)):
        p_el = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        p_el.append(pPr)

        # Alignment
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'),
               'center' if align == WD_ALIGN_PARAGRAPH.CENTER else
               'right'  if align == WD_ALIGN_PARAGRAPH.RIGHT  else 'left')
        pPr.append(jc)

        # Spacing
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:before'), str(space_before * 20))  # 20 = pt → twentieths
        spc.set(qn('w:after'),  str(space_after  * 20))
        pPr.append(spc)

        # Run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        if bold:
            rPr.append(OxmlElement('w:b'))
        if italic:
            rPr.append(OxmlElement('w:i'))
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(size_pt * 2))  # half-points
        rPr.append(sz)
        col = OxmlElement('w:color')
        col.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
        rPr.append(col)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p_el.append(r)
        return p_el

    paragraphs = []
    # Vertical spacing at top (≈8 line breaks)
    paragraphs.append(make_paragraph("", size_pt=14, space_after=120))

    # Book title — huge, centered
    paragraphs.append(make_paragraph(
        "KLİNİK PSİXİATRİYA",
        size_pt=44, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18,
    ))
    # Slogan — medium, italic
    paragraphs.append(make_paragraph(
        "Diaqnostika və terapiya standartları",
        size_pt=18, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    ))
    paragraphs.append(make_paragraph(
        "XBT-11 və DSM-5-TR əsasında klinik bələdçi",
        size_pt=14, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24, color=(85,85,85),
    ))

    # Big vertical gap (push author block toward the bottom)
    for _ in range(14):
        paragraphs.append(make_paragraph("", size_pt=12, space_after=0))

    # Author / city / year — bottom, centered
    paragraphs.append(make_paragraph(
        "Dr. Kənan Rəhimov",
        size_pt=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    ))
    paragraphs.append(make_paragraph(
        "Bakı · 2026",
        size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0, color=(85,85,85),
    ))

    # Page break after title page so TOC starts on new page
    pb_p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    pb_p.append(r)
    paragraphs.append(pb_p)

    return paragraphs


if __name__ == "__main__":
    main()
