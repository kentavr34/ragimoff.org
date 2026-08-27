#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_print_book.py — full print-ready DOCX for each language (az/ru/en/tr).

Unlike the old build_book.py (which only compiled the short chapter-overview
pages), this pulls in every disorder card's full content in chapter order,
builds a literal (non-field) page-numbered table of contents via a two-pass
LibreOffice-headless render, and writes one KLINIK_PSIXIATRIYA_<lang>.docx
per language into an output folder.

Usage:
    python build_print_book.py az
    python build_print_book.py ru
    python build_print_book.py all
"""
from __future__ import annotations
import re
import sys
import io
import json
import subprocess
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

ROOT = Path(__file__).parent.resolve()
SITE = ROOT / "klinik-psixiatriya"
BUILD_DIR = ROOT / "_build_print"
BUILD_DIR.mkdir(exist_ok=True)

SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")

CHAPTER_ORDER = [
    "01-6A0-neyroinkisaf.html", "02-6A2-sizofreniya-spektri.html",
    "03-6A4-katatoniya.html", "04-6A6-ehval-pozuntulari.html",
    "05-6B0-narahatliq.html", "06-6B2-okp.html", "07-6B4-stress.html",
    "08-6B6-dissosiativ.html", "09-6B8-qida-qebulu.html",
    "10-6C0-ifrazat.html", "11-6C2-bedensel-disstres.html",
    "12-6C4-madde-asililiq.html", "13-6C7-impuls-nezareti.html",
    "14-6C9-pozucu-davranis.html", "15-6D1-sexsiyyet.html",
    "16-6D3-parafilik.html", "17-6D5-faktitioz.html",
    "18-6D7-neyrokoqnitiv.html", "19-6E2-perinatal.html",
    "20-6E4-psixosomatik.html", "21-6E6-ikincili.html",
    "22-7AB-yuxu.html", "23-HA-cinsi-saglamliq.html",
]
APPENDIX_FILES = ["elave-acde.html", "elave-skalalar.html", "yekun.html"]

LANGS = {
    "az": dict(
        dirname="", title="KLİNİK PSİXİATRİYA",
        slogan="Diaqnostika və terapiya standartları",
        subslogan="XBT-11 və DSM-5-TR əsasında klinik bələdçi",
        author="Dr. Kənan Rəhimov", city_year="Bakı · 2026",
        toc_title="MÜNDƏRİCAT", preface_title="MÜQƏDDİMƏ",
        appendix_title="ƏLAVƏLƏR", abbr_title="QISALTMALAR",
        header_book="KLİNİK PSİXİATRİYA · Diaqnostika və terapiya standartları",
        icd_lbl="XBT-11", icd10_lbl="XBT-10", dsm_lbl="DSM-5-TR",
        page_word="səh.",
    ),
    "ru": dict(
        dirname="ru", title="КЛИНИЧЕСКАЯ ПСИХИАТРИЯ",
        slogan="Стандарты диагностики и терапии",
        subslogan="Клиническое руководство на основе МКБ-11 и DSM-5-TR",
        author="Д-р Кенан Рагимов", city_year="Баку · 2026",
        toc_title="СОДЕРЖАНИЕ", preface_title="ПРЕДИСЛОВИЕ",
        appendix_title="ПРИЛОЖЕНИЯ", abbr_title="СОКРАЩЕНИЯ",
        header_book="КЛИНИЧЕСКАЯ ПСИХИАТРИЯ · Стандарты диагностики и терапии",
        icd_lbl="МКБ-11", icd10_lbl="МКБ-10", dsm_lbl="DSM-5-TR",
        page_word="стр.",
    ),
    "en": dict(
        dirname="en", title="CLINICAL PSYCHIATRY",
        slogan="Diagnostic and treatment standards",
        subslogan="A clinical guide based on ICD-11 and DSM-5-TR",
        author="Dr. Kanan Ragimov", city_year="Baku · 2026",
        toc_title="CONTENTS", preface_title="PREFACE",
        appendix_title="APPENDICES", abbr_title="ABBREVIATIONS",
        header_book="CLINICAL PSYCHIATRY · Diagnostic and treatment standards",
        icd_lbl="ICD-11", icd10_lbl="ICD-10", dsm_lbl="DSM-5-TR",
        page_word="p.",
    ),
    "tr": dict(
        dirname="tr", title="KLİNİK PSİKİYATRİ",
        slogan="Tanı ve tedavi standartları",
        subslogan="ICD-11 ve DSM-5-TR temelli klinik rehber",
        author="Dr. Kenan Rahimov", city_year="Bakü · 2026",
        toc_title="İÇİNDEKİLER", preface_title="ÖNSÖZ",
        appendix_title="EKLER", abbr_title="KISALTMALAR",
        header_book="KLİNİK PSİKİYATRİ · Tanı ve tedavi standartları",
        icd_lbl="ICD-11", icd10_lbl="ICD-10", dsm_lbl="DSM-5-TR",
        page_word="s.",
    ),
}

MAIN_RE = re.compile(r'<main\b[^>]*>(.*?)</main>', re.DOTALL | re.IGNORECASE)
NAV_TREE_RE = re.compile(
    r'<div class="cls-tree" data-cls="icd">(.*?)</div>\s*<div class="cls-tree" data-cls="dsm"',
    re.DOTALL,
)
CHAPTER_RE = re.compile(
    r'<div class="nav-item nav-has-sub"><a href="([^"]+)" class="nav-link is-bolme"'
    r'[^>]*><span class="nav-code">([^<]*)</span><span>([^<]*)</span></a>.*?'
    r'<div class="nav-sub">(.*?)</div></div>',
    re.DOTALL,
)
SUB_RE = re.compile(
    r'<a href="([^"]+)" class="nav-sub-link" data-slug="([^"]+)">'
    r'<span class="sub-code">([^<]*)</span><span class="sub-name">([^<]*)</span></a>'
)


def lang_dir(lang: str) -> Path:
    d = LANGS[lang]["dirname"]
    return SITE / d if d else SITE


def get_chapter_map(lang: str) -> list[dict]:
    """Parse this language's own index.html nav for chapter+disorder names,
    keeping only the 23 real chapters (in CHAPTER_ORDER)."""
    t = (lang_dir(lang) / "index.html").read_text(encoding="utf-8")
    m = NAV_TREE_RE.search(t)
    frag = m.group(1)
    by_file = {}
    for cm in CHAPTER_RE.finditer(frag):
        href, code_range, name, subs_html = cm.groups()
        if href not in CHAPTER_ORDER:
            continue
        subs = []
        for sm in SUB_RE.finditer(subs_html):
            sub_href, slug, sub_code, sub_name = sm.groups()
            subs.append({"code": sub_code, "name": sub_name, "file": sub_href})
        by_file[href] = {"file": href, "code_range": code_range, "name": name, "disorders": subs}
    return [by_file[f] for f in CHAPTER_ORDER if f in by_file]


def extract_main(html: str) -> str:
    m = MAIN_RE.search(html)
    if not m:
        return ""
    body = m.group(1)
    body = re.sub(r'<!--\s*/?BOOK-SUPPLEMENT:[^>]*-->', '', body, flags=re.IGNORECASE)
    body = re.sub(r'<aside\b.*?</aside>', '', body, flags=re.DOTALL | re.IGNORECASE)
    body = re.sub(r'<script\b.*?</script>', '', body, flags=re.DOTALL | re.IGNORECASE)
    body = re.sub(r'<style\b.*?</style>', '', body, flags=re.DOTALL | re.IGNORECASE)
    body = re.sub(r'<nav\b.*?</nav>', '', body, flags=re.DOTALL | re.IGNORECASE)
    prev = None
    while prev != body:
        prev = body
        body = re.sub(r'<abbr\b[^<>]*>([^<]*)</abbr>', r'\1', body, flags=re.IGNORECASE)
    body = re.sub(r'</?abbr\b[^>]*>', '', body, flags=re.IGNORECASE)
    return body.strip()


def strip_links(html: str) -> str:
    return re.sub(r'<a\b[^>]*>(.*?)</a>', r'\1', html, flags=re.DOTALL | re.IGNORECASE)


def strip_chapter_web_nav(html: str) -> str:
    """Chapter overview pages carry a clickable disorder-shortcut grid
    (`<div class="chapter-menu">`) and a lead-in sentence pointing at it
    ("each on its own page — click below"). Both are pure web-navigation
    aids: with links stripped for print, the grid collapses into an
    unreadable run-on string, and the sentence stops making sense once the
    grid is gone. The disorders already get their own page-numbered TOC
    entries and page breaks, so drop both."""
    html = re.sub(r'<div class="chapter-menu">.*?</div>', '', html, flags=re.DOTALL)
    html = re.sub(r'<p class="chap-sub">.*?</p>', '', html, flags=re.DOTALL)
    return html


def fix_dh_table_linebreaks(html: str) -> str:
    """The disorder/chapter header table relies on CSS `display:block` to
    stack the label (XBT-11/XBT-10/DSM-5-TR) above its code — pandoc has no
    CSS, so without an explicit break they'd run together as 'XBT-116A00'."""
    return re.sub(
        r'(<span class="dh-lbl">[^<]*</span>)(<span class="dh-code">)',
        r'\1<br>\2', html,
    )


def strip_inline_bold(html: str) -> str:
    def strip_in(tag, text):
        pattern = re.compile(rf'(<{tag}\b[^>]*>)(.*?)(</{tag}>)', re.DOTALL | re.IGNORECASE)
        def repl(m):
            inner = re.sub(r'</?(?:strong|b)\b[^>]*>', '', m.group(2), flags=re.IGNORECASE)
            return m.group(1) + inner + m.group(3)
        return pattern.sub(repl, text)
    for tag in ("p", "li", "em"):
        html = strip_in(tag, html)
    return html


def shift_headings(html: str, delta: int) -> str:
    """Shift every <hN> by delta levels (clamped to 1..6)."""
    def repl(m):
        slash, lvl, attrs = m.group(1), int(m.group(2)), m.group(3)
        new_lvl = max(1, min(lvl + delta, 6))
        return f'<{slash}h{new_lvl}{attrs}>'
    return re.sub(r'<(/?)h([1-6])(\b[^>]*)>', repl, html)


DISORDER_H1_RE = re.compile(r'<h1[^>]*>(.*?)</h1>', re.DOTALL)
CHAPTER_H1_RE = re.compile(
    r'<h1 class="chap-h1"><span class="chap-range">(.*?)</span>'
    r'<span class="chap-title">(.*?)</span></h1>', re.DOTALL
)


def gather_disorder(lang: str, disorder: dict, marker_idx: int) -> tuple[str, str]:
    """Full <main> of one disorder card, headings shifted +1
    (h1 disorder-title -> h2, h2 section -> h3, h3 sub -> h4, h4 myth -> h5).
    The h1 is rewritten to 'CODE  NAME' for a proper printed/TOC heading and
    carries a ‡‡Mi‡‡ marker (injected here, BEFORE the shift, since the
    marker must land on whatever tag is still literally <h1> at this point —
    inject_markers() only looks for <h1>, and after shifting this becomes
    <h2>). Returns (body_html, heading_text) — heading_text is marker-free,
    safe to show verbatim in the TOC / use as the PDF-search anchor label.
    """
    path = lang_dir(lang) / disorder["file"]
    if not path.exists():
        print(f"  MISSING disorder file: {disorder['file']}")
        return "", ""
    html = path.read_text(encoding="utf-8", errors="ignore")
    body = extract_main(html)
    if not body:
        return "", ""
    heading_text = f'{disorder["code"]}  {_clean(disorder["name"])}'
    body, n = DISORDER_H1_RE.subn(f'<h1>‡‡M{marker_idx}‡‡{heading_text}</h1>', body, count=1)
    body = fix_dh_table_linebreaks(body)
    body = shift_headings(body, 1)
    body = strip_links(body)
    body = strip_inline_bold(body)
    return body, heading_text


def gather_chapter_intro(lang: str, chapter_file: str, marker_idx: int) -> tuple[str, str]:
    """Chapter overview <main> — its own h1 is the book's Heading 1 (chapter
    title), rewritten to 'RANGE  NAME' for a clean single-run heading, with
    a ‡‡Mi‡‡ marker prepended the same way as gather_disorder (chapter
    intros aren't shifted, but keeping the injection site consistent avoids
    a second, easy-to-forget code path).
    Returns (body_html, heading_text)."""
    path = lang_dir(lang) / chapter_file
    html = path.read_text(encoding="utf-8", errors="ignore")
    body = extract_main(html)
    body = strip_chapter_web_nav(body)
    m = CHAPTER_H1_RE.search(body)
    if m:
        code_range, name = _clean(m.group(1)), _clean(m.group(2))
        heading_text = f'{code_range}  {name}'
        body = CHAPTER_H1_RE.sub(f'<h1>‡‡M{marker_idx}‡‡{heading_text}</h1>', body, count=1)
    else:
        heading_text = ""
    body = strip_links(body)
    body = strip_inline_bold(body)
    return body, heading_text


def gather_frontback(lang: str, fname: str, shift: int = 0) -> str:
    """Preface / appendix / abbreviations page — own h1(s) stay as book
    Heading-1-level chapter markers; everything else shifted by `shift`."""
    path = lang_dir(lang) / fname
    if not path.exists():
        print(f"  MISSING front/back file: {fname}")
        return ""
    html = path.read_text(encoding="utf-8", errors="ignore")
    body = extract_main(html)
    if shift:
        body = shift_headings(body, shift)
    body = strip_links(body)
    body = strip_inline_bold(body)
    return body


ICD_RE = re.compile(r'^\s*([0-9][A-Z][0-9][0-9A-Z]?|HA[0-9]{2})\b')
H1_OPEN_RE = re.compile(r'(<h1[^>]*>)')


def inject_markers(html: str, start_idx: int) -> tuple[str, int]:
    """Prepend a literal, uncollapsible '‡‡M<i>‡‡' token right inside every
    <h1> in `html`, one index per h1 in document order, starting at
    start_idx. Pandoc's HTML reader collapses whitespace runs (a double
    space between code and name doesn't survive to the DOCX), so a plain
    marker character sequence — not a whitespace pattern — is the only
    reliable way later code has to tell 'this is a toc-tracked heading' and
    'this is toc entry N' apart from an incidental subsection heading in
    front/back-matter prose. Returns (new_html, count_of_h1s)."""
    count = 0
    def repl(m):
        nonlocal count
        marker = f'‡‡M{start_idx + count}‡‡'
        count += 1
        return m.group(1) + marker
    new_html = H1_OPEN_RE.sub(repl, html)
    return new_html, count


def build_book_parts(lang: str) -> tuple[list[str], list[dict]]:
    """Returns (html_parts, toc_entries).
    toc_entries: list of {level: 'chapter'|'disorder', text, anchor}
    `anchor` is the exact heading text to search for in the rendered PDF
    (identical to what actually lands as the <h1> text in that part);
    `text` is what gets printed in the TOC (may differ for disorders, which
    get a code prefix even though anchor already includes it here).
    """
    chapters = get_chapter_map(lang)
    parts = []
    toc = []

    def add_part_needing_injection(html: str, entries: list[dict]) -> None:
        """Preface/appendix/abbreviations: h1's are still literally <h1> at
        this point (never shifted), so inject_markers() can find and mark
        each one generically."""
        marked_html, n = inject_markers(html, len(toc))
        if n != len(entries):
            print(f"  WARNING: {n} <h1> found but {len(entries)} toc entries expected")
        parts.append(marked_html)
        toc.extend(entries)

    def add_part_premarked(html: str, entry: dict) -> None:
        """Chapter/disorder: gather_chapter_intro()/gather_disorder() already
        baked the ‡‡Mi‡‡ marker into their (possibly since-shifted) heading,
        using len(toc) as the index at call time — just keep both in sync."""
        parts.append(html)
        toc.append(entry)

    # Preface
    pref_html = gather_frontback(lang, "mugeddime.html")
    m = re.search(r'<h1[^>]*>(.*?)</h1>', pref_html)
    entries = []
    if m:
        text = _clean(m.group(1))
        entries.append({"level": "chapter", "text": text, "anchor": text})
    add_part_needing_injection(pref_html, entries)

    # Chapters + disorders
    for ch in chapters:
        intro, chapter_heading = gather_chapter_intro(lang, ch["file"], len(toc))
        if not chapter_heading:
            chapter_heading = f'{ch["code_range"]}  {ch["name"]}'
        add_part_premarked(intro, {"level": "chapter", "text": chapter_heading, "anchor": chapter_heading})
        for d in ch["disorders"]:
            dbody, dtext = gather_disorder(lang, d, len(toc))
            if not dbody:
                continue
            add_part_premarked(dbody, {"level": "disorder", "text": dtext, "anchor": dtext})

    # Appendices (each may have multiple own h1's)
    for fname in APPENDIX_FILES:
        body = gather_frontback(lang, fname)
        entries = [{"level": "chapter", "text": _clean(m.group(1)), "anchor": _clean(m.group(1))}
                   for m in re.finditer(r'<h1[^>]*>(.*?)</h1>', body)]
        add_part_needing_injection(body, entries)

    # Abbreviations glossary intentionally dropped from the print book —
    # owner's decision 2026-08-28: distracts, adds no clinical value,
    # trims ~20 pages. Kept on disk as abbreviatur.html/terminoloji-luget.html
    # (unlinked from site nav) purely as a local reference file.

    return parts, toc


def _clean(s: str) -> str:
    s = re.sub(r'<[^>]+>', '', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


# ─────────────────────────── DOCX conversion ────────────────────────────

def write_combined_html(parts: list[str], lang: str, out_path: Path) -> None:
    combined = "\n".join(parts)
    full = (
        "<!DOCTYPE html>\n<html><head><meta charset='utf-8'>"
        f"<title>{LANGS[lang]['title']}</title></head><body>\n"
        + combined + "\n</body></html>"
    )
    out_path.write_text(full, encoding="utf-8")


def convert_to_docx(html_path: Path, docx_path: Path) -> None:
    import pypandoc
    pypandoc.convert_file(
        str(html_path), "docx", outputfile=str(docx_path),
        extra_args=["--standalone"], format="html",
    )


def render_pdf(docx_path: Path, out_dir: Path) -> Path:
    subprocess.run(
        [str(SOFFICE), "--headless", "--convert-to", "pdf",
         "--outdir", str(out_dir), str(docx_path)],
        check=True, capture_output=True, timeout=600,
    )
    return out_dir / (docx_path.stem + ".pdf")


def detect_page_numbers(pdf_path: Path, toc: list[dict]) -> dict[int, int]:
    """Find the 1-based PDF page holding each heading's invisible per-index
    marker (‡‡M{i}‡‡, inserted by assemble_docx right inside the heading's
    own paragraph). Markers are unambiguous regardless of table-cell line
    wrapping or interleaving between adjacent columns — unlike the visible
    heading text itself, which cannot be reliably substring-matched once a
    disorder name wraps across lines next to its code column.
    Returns {toc_index: page_number}."""
    import pdfplumber
    pages_text = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            pages_text.append(re.sub(r'\s+', ' ', (page.extract_text() or '')))

    result = {}
    for i in range(len(toc)):
        marker = f'‡‡M{i}‡‡'
        found_page = None
        for pno, txt in enumerate(pages_text):
            if marker in txt:
                found_page = pno + 1
                break
        result[i] = found_page
    return result


# ─────────────────────────── Typography (TYPOGRAPHY.md) ────────────────

def _p(text, *, size_pt, bold=False, italic=False, align="center",
       space_before=0, space_after=0, color=(0, 0, 0), font="Times New Roman"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    p_el.append(pPr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), {"center": "center", "right": "right", "left": "left"}[align])
    pPr.append(jc)
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:before'), str(space_before * 20))
    spc.set(qn('w:after'), str(space_after * 20))
    pPr.append(spc)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rPr.append(sz)
    col = OxmlElement('w:color')
    col.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*color))
    rPr.append(col)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p_el.append(r)
    return p_el


def _page_break_p():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_el = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p_el.append(r)
    return p_el


def build_title_page(lang: str) -> list:
    L = LANGS[lang]
    parts = [
        _p("", size_pt=14, space_after=120),
        _p(L["title"], size_pt=44, bold=True, space_after=18),
        _p(L["slogan"], size_pt=18, italic=True, space_after=10),
        _p(L["subslogan"], size_pt=14, italic=True, space_after=24, color=(85, 85, 85)),
    ]
    parts += [_p("", size_pt=12, space_after=0) for _ in range(14)]
    parts += [
        _p(L["author"], size_pt=14, bold=True, space_after=4),
        _p(L["city_year"], size_pt=12, space_after=0, color=(85, 85, 85)),
        _page_break_p(),
    ]
    return parts


def build_toc(lang: str, toc: list[dict], page_numbers: dict[str, int] | None) -> list:
    """Literal (non-field) TOC: chapter lines flush-left bold, disorder lines
    indented with a dot-leader tab stop and the real page number typed as
    plain text at the right margin."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    L = LANGS[lang]
    out = [_p(L["toc_title"], size_pt=28, bold=True, space_after=24)]

    def toc_line(text, page_no, *, indent_cm, bold):
        p_el = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        p_el.append(pPr)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(indent_cm * 567)))
        pPr.append(ind)
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:before'), '0')
        spc.set(qn('w:after'), '60')
        pPr.append(spc)
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9639')  # ~17cm — right margin on A4 w/ 2cm margins
        tabs.append(tab)
        pPr.append(tabs)

        def add_run(txt, bold_run):
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            if bold_run:
                rPr.append(OxmlElement('w:b'))
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '22')
            rPr.append(sz)
            col = OxmlElement('w:color')
            col.set(qn('w:val'), '000000')
            rPr.append(col)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = txt
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p_el.append(r)

        add_run(text, bold)
        tab_r = OxmlElement('w:r')
        tab_t = OxmlElement('w:tab')
        tab_r.append(tab_t)
        p_el.append(tab_r)
        add_run(str(page_no) if page_no else "…", bold)
        return p_el

    for idx, e in enumerate(toc):
        pg = (page_numbers or {}).get(idx)
        if e["level"] == "chapter":
            out.append(toc_line(e["text"], pg, indent_cm=0, bold=True))
        else:
            out.append(toc_line(e["text"], pg, indent_cm=0.6, bold=False))
    out.append(_page_break_p())
    return out


def assemble_docx(docx_path: Path, lang: str, toc: list[dict],
                   page_numbers: dict[str, int] | None,
                   strip_markers: bool = True) -> None:
    """Post-process the raw pandoc DOCX: heading styles, page breaks, black
    text, headers/footers, table borders, and prepend title page + TOC.
    strip_markers=False keeps the ‡‡Mi‡‡ heading markers in the saved DOCX
    (pass 1, so render_pdf's PDF still has them for detect_page_numbers);
    the final pass strips them so they never reach the reader."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    L = LANGS[lang]
    doc = Document(str(docx_path))

    def style_heading(name, size_pt, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
        try:
            s = doc.styles[name]
        except KeyError:
            return
        s.font.name = "Times New Roman"
        s.font.size = Pt(size_pt)
        s.font.bold = True
        s.font.italic = italic
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.alignment = align
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    style_heading("Heading 1", 28, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_heading("Heading 2", 20, align=WD_ALIGN_PARAGRAPH.LEFT)
    style_heading("Heading 3", 14, align=WD_ALIGN_PARAGRAPH.LEFT)
    style_heading("Heading 4", 12, align=WD_ALIGN_PARAGRAPH.LEFT)
    style_heading("Heading 5", 11, align=WD_ALIGN_PARAGRAPH.LEFT, italic=True)
    for name in ("Heading 6", "Heading 7", "Heading 8", "Heading 9"):
        try:
            doc.styles[name].font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass
    try:
        hl = doc.styles["Hyperlink"]
        hl.font.color.rgb = RGBColor(0, 0, 0)
        hl.font.underline = False
    except KeyError:
        pass

    try:
        n = doc.styles["Normal"]
        n.font.name = "Times New Roman"
        n.font.size = Pt(11)
        n.paragraph_format.space_after = Pt(4)
        n.paragraph_format.space_before = Pt(0)
        n.paragraph_format.line_spacing = 1.25
        n.paragraph_format.first_line_indent = Cm(0.5)
        n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        pass

    for sty_name in ("First Paragraph", "FirstParagraph"):
        try:
            fp = doc.styles[sty_name]
            fp.font.name = "Times New Roman"
            fp.font.size = Pt(11)
            fp.paragraph_format.first_line_indent = Cm(0)
            fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass

    # ── Page size + margins: A4, ~2cm margins (matches TOC tab-stop pos) ──
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ── Headers/footers ──
    settings_el = doc.settings.element
    if settings_el.find(qn('w:evenAndOddHeaders')) is None:
        settings_el.append(OxmlElement('w:evenAndOddHeaders'))
    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.first_page_header.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        eh = section.even_page_header
        eh.is_linked_to_previous = False
        ehp = eh.paragraphs[0]
        ehp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = ehp.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.italic = True
        rPr_xml = (
            '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            '<w:sz w:val="18"/><w:i/></w:rPr>'
        )
        for xml_frag in (
            f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{rPr_xml}<w:fldChar w:fldCharType="begin"/></w:r>',
            f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{rPr_xml}<w:instrText xml:space="preserve"> STYLEREF "Heading 1" \\* MERGEFORMAT </w:instrText></w:r>',
            f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{rPr_xml}<w:fldChar w:fldCharType="end"/></w:r>',
        ):
            from docx.oxml import parse_xml
            ehp._p.append(parse_xml(xml_frag))
        oh = section.header
        oh.is_linked_to_previous = False
        ohp = oh.paragraphs[0]
        ohp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = ohp.add_run(L["header_book"])
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(9)
        r2.italic = True
        for foot, align in [(section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
                             (section.footer, WD_ALIGN_PARAGRAPH.RIGHT)]:
            foot.is_linked_to_previous = False
            fp_ = foot.paragraphs[0]
            fp_.alignment = align
            from docx.oxml import parse_xml
            NSW = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            fp_._p.append(parse_xml(f'<w:r {NSW}><w:fldChar w:fldCharType="begin"/></w:r>'))
            fp_._p.append(parse_xml(f'<w:r {NSW}><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'))
            fp_._p.append(parse_xml(f'<w:r {NSW}><w:fldChar w:fldCharType="end"/></w:r>'))

    # ── Table cell indent/borders ──
    def kill_indent(par):
        pPr = par._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:left'), '0')

    def set_cell_border(cell):
        tc_pr = cell._tc.get_or_add_tcPr()
        tcBorders = tc_pr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tc_pr.append(tcBorders)
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:color'), '666666')
            tcBorders.append(b)

    for tbl in doc.tables:
        tbl.autofit = True
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_border(cell)
                for par in cell.paragraphs:
                    kill_indent(par)

    # ── Page-break-before rules, driven by the ‡‡M<i>‡‡ markers baked into
    #    each toc-tracked <h1> back in build_book_parts(). Marker index i
    #    maps 1:1 onto toc[i], so toc[i]["level"] says chapter vs disorder
    #    without re-deriving it from style/text heuristics (Heading 2 is
    #    also used, un-shifted, by ordinary front/back-matter subsections
    #    like mugeddime.html's "Oxucuya müraciət" — those must NOT break). ──
    MARKER_RE = re.compile(r'‡‡M(\d+)‡‡')
    H1_STYLES = {'Heading1', 'Heading 1', 'heading 1'}
    H2_STYLES = {'Heading2', 'Heading 2', 'heading 2'}
    SUB_HEADINGS = {'Heading3', 'Heading 3', 'heading 3', 'Heading4', 'Heading 4',
                    'heading 4', 'Heading5', 'Heading 5', 'heading 5'}
    body = doc.element.body

    def text_of(p):
        return "".join(t.text or "" for t in p.iter(qn('w:t')))

    def add_pbb(p):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)
        if pPr.find(qn('w:pageBreakBefore')) is None:
            pPr.insert(0, OxmlElement('w:pageBreakBefore'))

    def remove_pbb(p):
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            return
        pbb = pPr.find(qn('w:pageBreakBefore'))
        if pbb is not None:
            pPr.remove(pbb)

    def strip_marker_text(p, marker_str):
        """Remove the literal '‡‡Mi‡‡' run text without disturbing any
        other runs in the paragraph (it's always its own whole run — see
        inject_markers — but be defensive and substring-replace anyway)."""
        for t in p.iter(qn('w:t')):
            if t.text and marker_str in t.text:
                t.text = t.text.replace(marker_str, '')

    # Recursive: disorder headings (Heading 2) live inside the header-table
    # cell, not as top-level body paragraphs — plain findall() (direct
    # children only) would silently skip every one of them.
    paragraphs = list(body.iter(qn('w:p')))
    found_markers = 0
    for p in paragraphs:
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            continue
        sty = pStyle.get(qn('w:val')) or ''
        if sty in SUB_HEADINGS:
            remove_pbb(p)
            continue
        if sty not in H1_STYLES and sty not in H2_STYLES:
            continue
        m = MARKER_RE.search(text_of(p))
        if not m:
            # un-shifted front/back-matter subsection, e.g. Heading 2
            # "Oxucuya müraciət" — not toc-tracked, never breaks
            remove_pbb(p)
            continue
        idx = int(m.group(1))
        if strip_markers:
            strip_marker_text(p, m.group(0))
        found_markers += 1
        if idx < len(toc) and toc[idx]["level"] == "disorder" and sty not in H1_STYLES:
            add_pbb(p)
        elif sty in H1_STYLES:
            add_pbb(p)
        else:
            remove_pbb(p)

    if found_markers != len(toc):
        print(f"  WARNING: {found_markers} markers found in body vs {len(toc)} TOC entries")

    # first heading: no leading blank page
    for p in paragraphs:
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            continue
        if (pStyle.get(qn('w:val')) or '') in H1_STYLES | H2_STYLES:
            remove_pbb(p)
            break

    # ── Force black on every text run (except title-page grey) ──
    GREY = '555555'
    def sweep_color(par_el):
        for r in par_el.findall(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                continue
            col = rPr.find(qn('w:color'))
            if col is None:
                continue
            if (col.get(qn('w:val')) or '').upper() == GREY.upper():
                continue
            col.set(qn('w:val'), '000000')
            for attr in ('w:themeColor', 'w:themeShade', 'w:themeTint'):
                if col.get(qn(attr)) is not None:
                    del col.attrib[qn(attr)]
    for p in body.findall(qn('w:p')):
        sweep_color(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    sweep_color(par._p)

    # ── Prepend title page + TOC ──
    title_paragraphs = build_title_page(lang)
    toc_paragraphs = build_toc(lang, toc, page_numbers)
    for tp in reversed(toc_paragraphs):
        body.insert(0, tp)
    for tp in reversed(title_paragraphs):
        body.insert(0, tp)

    doc.save(str(docx_path))
    _purge_theme_blue(docx_path)


def _purge_theme_blue(path: Path):
    import zipfile
    GREY = '555555'
    tmp = Path(str(path) + ".tmp.zip")

    def fix_styles(xml: str) -> str:
        def repl(m):
            tag = m.group(0)
            mv = re.search(r'w:val="([0-9A-Fa-f]+)"', tag)
            val = (mv.group(1) if mv else '').upper()
            new_val = val if val == GREY.upper() else '000000'
            self_closing = tag.endswith('/>')
            return f'<w:color w:val="{new_val}"' + ('/>' if self_closing else '>')
        return re.sub(r'<w:color\b[^/>]*?/?>', repl, xml)

    def fix_theme(xml: str) -> str:
        return re.sub(r'<a:srgbClr\s+val="[0-9A-Fa-f]+"\s*/>', '<a:srgbClr val="000000"/>', xml)

    with zipfile.ZipFile(str(path), 'r') as zin, \
         zipfile.ZipFile(str(tmp), 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == 'word/styles.xml':
                data = fix_styles(data.decode('utf-8')).encode('utf-8')
            elif item == 'word/theme/theme1.xml':
                data = fix_theme(data.decode('utf-8')).encode('utf-8')
            zout.writestr(item, data)
    path.unlink()
    tmp.rename(path)


OUT_DIR = Path(r"C:\Users\KENAN\Desktop\Клиническая психиатрия")


def build_one(lang: str) -> Path:
    print(f"=== {lang} ===")
    work = BUILD_DIR / lang
    work.mkdir(exist_ok=True)

    print("  gathering content...")
    parts, toc = build_book_parts(lang)
    print(f"  {len(parts)} parts, {len(toc)} TOC entries")

    html_path = work / "combined.html"
    write_combined_html(parts, lang, html_path)

    docx_pass1 = work / "pass1.docx"
    print("  pandoc -> docx (pass 1)...")
    convert_to_docx(html_path, docx_pass1)
    assemble_docx(docx_pass1, lang, toc, page_numbers=None, strip_markers=False)

    print("  soffice -> pdf...")
    pdf_path = render_pdf(docx_pass1, work)

    print("  detecting page numbers...")
    page_numbers = detect_page_numbers(pdf_path, toc)
    missing = [i for i, v in page_numbers.items() if v is None]
    if missing:
        print(f"  WARNING: {len(missing)} headings not found in PDF:")
        for i in missing[:10]:
            print("    -", toc[i]["anchor"][:70])

    docx_pass2 = work / "pass2.docx"
    print("  pandoc -> docx (pass 2, final)...")
    convert_to_docx(html_path, docx_pass2)
    assemble_docx(docx_pass2, lang, toc, page_numbers=page_numbers)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    final_name = {
        "az": "Klinik_Psixiatriya.docx", "ru": "Klinicheskaya_Psikhiatriya.docx",
        "en": "Clinical_Psychiatry.docx", "tr": "Klinik_Psikiyatri.docx",
    }[lang]
    final_path = OUT_DIR / final_name
    import shutil
    shutil.copy(docx_pass2, final_path)
    print(f"  DONE: {final_path}")
    return final_path


if __name__ == "__main__":
    lang = sys.argv[1] if len(sys.argv) > 1 else "az"
    if lang == "all":
        for l in LANGS:
            build_one(l)
    elif lang == "map":
        for l in LANGS:
            chapters = get_chapter_map(l)
            print(l, "chapters:", len(chapters), "disorders:",
                  sum(len(c["disorders"]) for c in chapters))
    else:
        build_one(lang)
